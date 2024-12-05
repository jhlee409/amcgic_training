import streamlit as st
from PIL import Image
import docx
import io
from io import BytesIO
import base64
import os
import firebase_admin
from firebase_admin import credentials, storage

# Set page to wide mode
st.set_page_config(page_title="EGD_Dx", layout="wide")

if st.session_state.get('logged_in'):

    # Check if Firebase app has already been initialized
    if not firebase_admin._apps:
        # Streamlit Secrets에서 Firebase 설정 정보 로드
        cred = credentials.Certificate({
            "type": "service_account",
            "project_id": st.secrets["project_id"],
            "private_key_id": st.secrets["private_key_id"],
            "private_key": st.secrets["private_key"].replace('\\n', '\n'),
            "client_email": st.secrets["client_email"],
            "client_id": st.secrets["client_id"],
            "auth_uri": st.secrets["auth_uri"],
            "token_uri": st.secrets["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["client_x509_cert_url"],
            "universe_domain": st.secrets["universe_domain"]
        })
        firebase_admin.initialize_app(cred)

    # Display Form Title_
    st.subheader("EGD_Dx_training")
    with st.expander(" 필독!!! 먼저 여기를 눌러 사용방법을 확인하세요."):
        st.write("- 가장 먼저 왼쪽 sidebar에서 F1용인지 F2용인지를 선택합니다.")
        st.write("- 그 아래에서 EGD 사진을 선택해서 업로드 하세요. 다음은 질문이 나옵니다. 잘 생각해 보고 결론이 났으면 왼쪽 sidebar에서 '진행' 버튼을 눌러 답과 설명을 보세요.")
        st.write("- 정리: 가 나오면 증례가 종결된 것입니다. 다음 증례를 위해 초기화를 하거나, 로그아웃을 하세요.")
        
    # Firebase에서 이미지를 다운로드하고 PIL 이미지 객체로 열기
    def download_and_open_image(bucket_name, file_path):
        bucket = storage.bucket(bucket_name)
        blob = bucket.blob(file_path)
        # BytesIO 객체를 사용하여 메모리에서 이미지를 직접 열기
        image_stream = io.BytesIO()
        blob.download_to_file(image_stream)
        image_stream.seek(0)
        return Image.open(image_stream)
        
    # Function to list files in a specific directory in Firebase Storage
    def png_list_files(bucket_name, directory):
        bucket = storage.bucket(bucket_name)
        blobs = bucket.list_blobs(prefix=directory)
        file_names = []
        for blob in blobs:
            # Extracting file name from the path and adding to the list
            file_name = blob.name[len(directory):]  # Remove directory path from file name
            if file_name:  # Check to avoid adding empty strings (in case of directories)
                file_names.append(file_name)
        return file_names
    
    # F1 or F2 selection
    folder_selection = st.sidebar.radio("Select Folder", ["초기화", "F1", "F2", "working"])

    if folder_selection == "초기화":
        directory_images = "EGD_Dx_training/Default/images/"
        directory_instructions = "EGD_Dx_training/Default/instructions/"

    elif folder_selection == "F1":
        directory_images = "EGD_Dx_training/F1/images/"
        directory_instructions = "EGD_Dx_training/F1/instructions/"
    elif folder_selection == "F2":
        directory_images = "EGD_Dx_training/F2/images/"
        directory_instructions = "EGD_Dx_training/F2/instructions/"
    else:
        directory_images = "EGD_Dx_training/working/images/"
        directory_instructions = "EGD_Dx_training/working/instructions/"

    st.sidebar.divider()

    # List and select PNG files
    file_list_images = png_list_files('amcgi-bulletin.appspot.com', directory_images)
    selected_image_file = st.sidebar.selectbox(f"EGD 사진을 선택하세요.", file_list_images)

    if selected_image_file:
        selected_image_path = directory_images + selected_image_file
        image = download_and_open_image('amcgi-bulletin.appspot.com', selected_image_path)
        
        # Open the image to check its dimensions
        width, height = image.size
        
        # Determine the display width based on the width-height ratio
        display_width = 1400 if width >= 1.6 * height else 700
        
        # Convert the image to base64
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        
        # Create the HTML code with oncontextmenu attribute
        html_code = f'<img src="data:image/png;base64,{img_str}" alt="Image" width="{display_width}" oncontextmenu="return false;" style="pointer-events: none;">'
        
        # Display the image using HTML code
        st.markdown(html_code, unsafe_allow_html=True)

        # Find the corresponding _1.docx file
        docx_file_name_1 = os.path.splitext(selected_image_file)[0] + '_1.docx'
        docx_file_path_1 = directory_instructions + docx_file_name_1
        
        # Download and read the contents of the _1.docx file
        bucket = storage.bucket('amcgi-bulletin.appspot.com')
        blob = bucket.blob(docx_file_path_1)
        docx_stream_1 = io.BytesIO()
        blob.download_to_file(docx_stream_1)
        docx_stream_1.seek(0)
        
        doc_1 = docx.Document(docx_stream_1)
        docx_content_1 = '\n'.join([paragraph.text for paragraph in doc_1.paragraphs])
        
        st.markdown(docx_content_1)

        st.divider()

        # Find the corresponding _2.docx file
        docx_file_name_2 = os.path.splitext(selected_image_file)[0] + '_2.docx'
        docx_file_path_2 = directory_instructions + docx_file_name_2
        
        # Download and read the contents of the _2.docx file
        blob = bucket.blob(docx_file_path_2)
        docx_stream_2 = io.BytesIO()
        blob.download_to_file(docx_stream_2)
        docx_stream_2.seek(0)
        
        doc_2 = docx.Document(docx_stream_2)
        docx_content_2 = '\n'.join([paragraph.text for paragraph in doc_2.paragraphs])
        
        if st.sidebar.button('진행'):
            st.markdown(docx_content_2)  # Show the content of _2.docx

    st.sidebar.divider()

    # 로그아웃 버튼 생성
    if st.sidebar.button('로그아웃'):
        st.session_state.logged_in = False
        st.rerun()  # 페이지를 새로고침하여 로그인 화면으로 돌아감

else:
    # 로그인이 되지 않은 경우, 로그인 페이지로 리디렉션 또는 메시지 표시
    st.error("로그인이 필요합니다.") 